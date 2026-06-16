from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Docs" / "DemoProject_Interview_Project_Document.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def style_para(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_para(p, after=3)
    r = p.add_run("DemoProject 项目说明文档")
    set_font(r, size=24, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    style_para(p, after=12)
    r = p.add_run("Unreal Engine 5.5 C++ Gameplay Demo | 独立完成 | 面试展示版")
    set_font(r, size=11, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    style_para(p)
    r = p.add_run(text)
    set_font(r, size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_para(p, after=4, line=1.167)
    r = p.add_run(text)
    set_font(r, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    style_para(p, after=4, line=1.167)
    r = p.add_run(text)
    set_font(r, size=11)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    style_para(p, after=3)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    style_para(p2, after=0)
    r2 = p2.add_run(text)
    set_font(r2, size=10.5)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_FILL)
        p = hdr[i].paragraphs[0]
        style_para(p, after=0)
        r = p.add_run(h)
        set_font(r, size=10.5, bold=True, color=RGBColor(0, 0, 0))
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            style_para(p, after=0, line=1.10)
            r = p.add_run(value)
            set_font(r, size=10)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def setup_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def build():
    doc = Document()
    setup_styles(doc)
    add_title(doc)

    add_callout(
        doc,
        "项目定位",
        "DemoProject 是一个基于 Unreal Engine 5.5 和 C++ 的动作 RPG / Top-down Gameplay Demo，重点展示 Gameplay Ability System、战斗技能、敌人 AI、背包、装备、拾取交互和 UMG 界面系统。项目的系统搭建、功能逻辑、C++ 实现、模块集成与调试均由我独立完成。",
    )

    add_table(
        doc,
        ["项目项", "说明"],
        [
            ["项目名称", "DemoProject"],
            ["开发引擎", "Unreal Engine 5.5"],
            ["主要语言", "C++，结合 Unreal Editor / UMG / Gameplay Tags 配置"],
            ["项目类型", "动作战斗、技能释放、敌人 AI、背包装备系统综合 Demo"],
            ["个人角色", "独立开发者，负责系统设计、编码实现、功能集成、调试与 GitHub 项目整理"],
            ["GitHub", "https://github.com/asadasdweq/DemoProject"],
        ],
        [2100, 7260],
    )

    add_heading(doc, "1. 项目概述")
    add_body(
        doc,
        "该项目围绕一个可展示的 UE5 Gameplay Demo 构建，核心目标是把角色战斗、技能释放、敌人目标选择、物品拾取、背包网格、装备穿戴、属性修改和界面反馈串成一条完整的游戏流程。项目不是单一功能练习，而是多个游戏系统之间的组合实现，适合作为面试中展示 C++ Gameplay 能力和系统拆分能力的作品。",
    )
    add_bullet(doc, "角色可以移动、拾取物品、释放技能、与敌人战斗。")
    add_bullet(doc, "技能系统基于 Gameplay Ability System，实现技能标签、冷却、伤害、属性和目标处理。")
    add_bullet(doc, "背包系统支持空间网格、堆叠、拆分、丢弃、消耗、装备和 UI 交互。")
    add_bullet(doc, "敌人 AI 使用 Behavior Tree 相关服务和任务完成目标选择与攻击行为。")
    add_bullet(doc, "UI 层包含 HUD、生命/魔法、技能显示、敌人血条、伤害文字、拾取提示和背包界面。")

    add_heading(doc, "2. 我的职责")
    add_body(
        doc,
        "整个 Demo 的功能系统由我独立完成，包括工程搭建、C++ Gameplay 框架、GAS 技能体系、属性与伤害逻辑、AI 行为、背包装备系统、UMG 界面交互、拾取高亮、网络复制相关处理、调试和 GitHub 仓库整理。面试中可以重点强调：这是一个独立完成的完整项目，而不是只参与其中某个模块。",
    )
    add_table(
        doc,
        ["职责方向", "具体内容"],
        [
            ["系统架构", "按照角色、技能、属性、AI、背包、装备、UI、交互等边界拆分模块，保证功能之间通过组件、接口、事件和 Gameplay Tag 解耦。"],
            ["C++ Gameplay", "实现角色初始化、ASC 绑定、属性初始化、技能授予、伤害计算、拾取检测、装备生成、背包操作等核心逻辑。"],
            ["UI 与交互", "实现背包网格、物品拖拽、堆叠拆分、右键菜单、拾取提示、HUD、敌人血条和伤害文字等界面反馈。"],
            ["网络意识", "在背包、装备、角色状态等系统中使用 Server RPC、Multicast、属性复制、FastArray 和 Replicated SubObject，保证逻辑以服务器为准。"],
            ["工程整理", "配置 Unreal 项目 Git 忽略规则，提交源码与资源到 GitHub，避免上传 Binaries、Intermediate、Saved 等生成目录。"],
        ],
        [2100, 7260],
    )

    add_heading(doc, "3. 核心功能模块")
    add_heading(doc, "3.1 Gameplay Ability System 战斗技能", level=2)
    add_body(
        doc,
        "项目使用 Gameplay Ability System 作为战斗系统基础。PlayerState 持有 AbilitySystemComponent 和 AttributeSet，并将 ASC 设置为 Mixed Replication Mode。角色在 PossessedBy / OnRep_PlayerState 阶段初始化 Ability Actor Info，随后初始化默认属性、授予初始技能并绑定 HUD。",
    )
    add_bullet(doc, "实现 FireBolt、Lightning、Electrocute、ChargeRock、MeleeAttack、ProjectileAbility 等能力类。")
    add_bullet(doc, "通过 Gameplay Tags 管理输入、技能、冷却、事件、状态、属性和伤害类型。")
    add_bullet(doc, "使用 SetByCaller Magnitude 传递不同伤害类型的数值。")
    add_bullet(doc, "实现 TargetDataUnderMouse、WaitCooldownChanged 等 AbilityTask / AsyncTask，支撑目标选择和冷却 UI。")

    add_heading(doc, "3.2 属性与伤害计算", level=2)
    add_body(
        doc,
        "属性系统分为 Vital、Primary、Secondary 等层级，包括 Health、Mana、Power、Intelligence、Vigor、Resilience、Armor、MagicalResistance、PhysicalDamage 和 MagicalDamage。伤害计算使用自定义 Execution Calculation：捕获源和目标属性，根据 Gameplay Tag 映射伤害类型与抗性，再将最终伤害写入 IncomingDamage。",
    )
    add_bullet(doc, "物理伤害通过 Armor 抵消，魔法伤害通过 MagicalResistance 抵消。")
    add_bullet(doc, "魔法伤害可叠加 Intelligence 带来的额外伤害。")
    add_bullet(doc, "通过 DamageTypeToResistance 映射，让后续扩展新的伤害类型更直接。")
    add_bullet(doc, "伤害流程与 GameplayEffectSpec、Captured Tags、Attribute Capture Defs 结合，符合 GAS 的标准扩展方式。")

    add_heading(doc, "3.3 背包与物品系统", level=2)
    add_body(
        doc,
        "背包系统是项目中复杂度最高的模块之一。InventoryComponent 负责物品添加、服务器校验、堆叠处理、丢弃、消耗、装备事件派发和 UI 构建。底层使用 FastArraySerializer 保存背包条目，并使用 Registered SubObject List 复制 UInv_InventoryItem 子对象。")
    add_bullet(doc, "拾取时进行服务器距离校验，避免客户端直接远程添加物品。")
    add_bullet(doc, "支持可堆叠物品的补满、剩余数量回写、部分拾取、部分丢弃。")
    add_bullet(doc, "支持消耗品，通过 Consumable Fragment 触发 GameplayEffect。")
    add_bullet(doc, "支持物品掉落，将当前物品 Manifest 重新生成 Pickup Actor。")

    add_heading(doc, "3.4 数据驱动的 Item Fragment 设计", level=2)
    add_body(
        doc,
        "物品数据使用 FInv_ItemManifest 加 TInstancedStruct<FInv_ItemFragment> 的组合方式。不同物品通过 Fragment 叠加能力，例如网格尺寸、图标、文本描述、可堆叠、可消耗、可装备、属性修改等。这样可以避免为每种物品写大量子类，同时方便在编辑器中配置不同物品行为。",
    )
    add_bullet(doc, "Grid Fragment 控制物品占用格子大小和 padding。")
    add_bullet(doc, "Image / Text / LabeledNumber Fragment 驱动物品 UI 展示。")
    add_bullet(doc, "Stackable Fragment 控制最大堆叠和当前数量。")
    add_bullet(doc, "Consumable / Equipment Fragment 分别处理消耗效果和装备效果。")
    add_bullet(doc, "Equip Modifier 可以通过 GameplayEffect 修改 Intelligence、Vigor、Resilience 等属性。")

    add_heading(doc, "3.5 空间网格背包 UI", level=2)
    add_body(
        doc,
        "背包界面采用空间网格设计，支持不同尺寸的物品占据多个格子。UI 逻辑会根据鼠标位置计算当前格子坐标、象限、起始放置位置、占用状态和可放置区域，并通过高亮、灰显、拖拽图标等方式给玩家反馈。")
    add_bullet(doc, "实现坐标与一维索引转换，使用 ForEach2D 遍历物品占用区域。")
    add_bullet(doc, "实现物品拾起、放下、交换、堆叠合并、拆分、拖出丢弃。")
    add_bullet(doc, "右键菜单支持 Split、Drop、Consume 等操作。")
    add_bullet(doc, "通过 SlotAvailabilityResult 统一表达空间、堆叠和剩余数量的判断结果。")

    add_heading(doc, "3.6 装备系统", level=2)
    add_body(
        doc,
        "装备系统监听背包组件的 OnItemEquipped / OnItemUnequipped 事件，读取物品 Manifest 中的 Equipment Fragment，生成装备 Actor 并附加到角色骨骼 Socket。装备时可以触发属性 GameplayEffect，卸下时移除或反向处理对应效果。")
    add_bullet(doc, "装备 Actor 与物品 Fragment 解耦，装备类型通过 Gameplay Tag 匹配。")
    add_bullet(doc, "武器装备后同步角色 HasWeapon 状态和武器 Mesh。")
    add_bullet(doc, "角色装备 Mesh 使用 Replicated 属性同步，保证客户端表现一致。")

    add_heading(doc, "3.7 AI、交互与 UI 反馈", level=2)
    add_body(
        doc,
        "敌人 AI 使用 AIController、Behavior Tree Service 和 Attack Task。BTService_FindNearestPlayer 根据 Player / Enemy Tag 查找最近目标，并将目标 Actor 与距离写入 Blackboard。交互层通过 Interface 实现拾取物高亮、敌人接口、战斗接口等能力，降低模块之间的直接依赖。")
    add_bullet(doc, "拾取范围使用 SphereComponent 检测，进入范围显示拾取提示并高亮物体。")
    add_bullet(doc, "HUD 使用 WidgetController 绑定属性和能力信息。")
    add_bullet(doc, "敌人血条、伤害文字、技能球、生命/魔法显示为战斗过程提供即时反馈。")

    add_heading(doc, "4. 核心技术亮点")
    add_table(
        doc,
        ["亮点", "技术实现", "面试可讲点"],
        [
            ["GAS 技能架构", "使用 ASC、AttributeSet、GameplayAbility、GameplayEffect、GameplayTag、AbilityTask 组织战斗逻辑。", "能说明 UE 中可扩展技能系统如何拆分输入、冷却、伤害、目标和 UI。"],
            ["自定义伤害计算", "通过 ExecCalc_Damage 捕获源/目标属性，按伤害类型和抗性映射计算最终 IncomingDamage。", "能讲清楚 SetByCaller、Attribute Capture、Tag 映射和抗性计算流程。"],
            ["数据驱动物品", "使用 ItemManifest + TInstancedStruct Fragment 组合物品行为。", "能解释为什么用组合代替大量继承，以及如何扩展新物品类型。"],
            ["复制背包系统", "InventoryComponent 使用 FastArraySerializer、Server RPC、Replicated SubObject 和 MarkItemDirty。", "能体现网络同步意识，知道哪些逻辑应该在服务器上处理。"],
            ["空间网格 UI", "实现格子坐标、象限、占用区域、堆叠合并、交换、拆分、灰显和高亮。", "能展示复杂 UI 交互和数据状态同步能力。"],
            ["装备与属性联动", "装备 Fragment 生成附加 Actor，并通过 GameplayEffect 修改角色属性。", "能说明装备系统如何和 GAS 属性系统连接。"],
            ["模块解耦", "通过 Gameplay Tags、Interface、Component、Delegate 分离角色、物品、UI、AI 和技能。", "能体现项目不是堆代码，而是有系统边界和扩展思路。"],
            ["工程完整度", "项目包含源码、资源、README、Git 忽略规则，并已推送到 GitHub。", "面试官可以直接查看工程结构和提交状态。"],
        ],
        [1500, 3100, 4760],
    )

    add_heading(doc, "5. 开发难点与解决思路")
    add_bullet(doc, "背包同步难点：物品既有列表关系，又有 UObject 子对象状态。解决方式是 FastArray 负责列表变更，Registered SubObject 负责物品对象复制，并在数量变化时主动 MarkItemDirty。")
    add_bullet(doc, "堆叠物品难点：拾取、拆分、丢弃和消耗都可能产生部分数量。解决方式是统一维护 StackCount、Remainder 和 SlotAvailabilityResult，避免 UI 与服务器逻辑各算一套。")
    add_bullet(doc, "空间格子难点：不同尺寸物品需要判断边界、占用、重叠、上左角索引和鼠标象限。解决方式是封装坐标/索引转换和 ForEach2D 遍历。")
    add_bullet(doc, "GAS 调试难点：伤害来源、属性捕获、Tag 和 GameplayEffect 之间链路较长。解决方式是把伤害类型、抗性和属性 Tag 集中管理，保证数据流可以追踪。")
    add_bullet(doc, "系统耦合难点：装备、背包、角色 Mesh、属性和 UI 都有关联。解决方式是通过 Fragment、Delegate、Component 和 Gameplay Tag 建立可扩展连接。")

    add_heading(doc, "6. 面试讲解建议")
    add_body(doc, "可以用 1 到 2 分钟先概括项目，然后按模块深入。建议不要一开始就讲所有代码，而是先讲项目目标和系统划分，再选两个最有价值的模块展开：GAS 战斗和背包装备。")
    add_number(doc, "开场：这是我独立完成的 UE5.5 C++ Gameplay Demo，重点展示 GAS 战斗、敌人 AI、背包、装备和 UMG 交互。")
    add_number(doc, "重点一：GAS 技能系统，包括 ASC 初始化、属性、技能授予、GameplayTag、冷却和自定义伤害计算。")
    add_number(doc, "重点二：背包装备系统，包括 FastArray 网络同步、Fragment 数据驱动、空间网格 UI、堆叠拆分和装备属性联动。")
    add_number(doc, "重点三：项目完整度，包括 GitHub 仓库、README、可运行 Demo 视频和打包版本。")

    add_heading(doc, "7. Demo 视频展示顺序")
    add_bullet(doc, "角色移动与镜头视角。")
    add_bullet(doc, "展示 FireBolt、Lightning、Electrocute、ChargeRock 等技能效果。")
    add_bullet(doc, "展示敌人追踪、攻击、受击、血条和伤害数字。")
    add_bullet(doc, "展示拾取物品、高亮提示和拾取消息。")
    add_bullet(doc, "打开背包，展示物品占格、拖拽、交换、堆叠、拆分、丢弃、消耗。")
    add_bullet(doc, "展示装备武器或装备类物品后角色状态/属性/外观变化。")
    add_bullet(doc, "最后展示 GitHub 仓库或打包运行画面，强调项目完整性。")

    add_heading(doc, "8. 可继续补充的材料")
    add_bullet(doc, "项目截图：主角战斗、技能释放、背包界面、装备界面、敌人血条。")
    add_bullet(doc, "演示视频链接：建议 1 到 3 分钟，上传到 B 站、YouTube、网盘或 GitHub Release 说明中。")
    add_bullet(doc, "打包版本：DemoProject_Windows.zip，可放 GitHub Release 或网盘。")
    add_bullet(doc, "README 优化：把视频链接、截图、操作方式和个人职责补充到 GitHub 首页。")

    add_heading(doc, "9. 简历/面试可用表述")
    add_callout(
        doc,
        "一句话版本",
        "独立完成基于 Unreal Engine 5.5 和 C++ 的动作 RPG Gameplay Demo，实现 GAS 技能战斗、属性伤害、敌人 AI、空间网格背包、装备属性联动、拾取交互和 UMG 界面，并整理 GitHub 仓库用于作品展示。",
    )
    add_callout(
        doc,
        "技术亮点版本",
        "项目中使用 Gameplay Ability System 构建可扩展技能框架，基于 Gameplay Tag 和自定义 Execution Calculation 处理伤害类型、抗性和属性加成；背包系统使用 FastArraySerializer 与 Replicated SubObject 完成网络同步，并通过 ItemManifest + Fragment 组合实现数据驱动物品、消耗品、装备和空间网格 UI。",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
